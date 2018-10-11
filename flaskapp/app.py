from flask import Flask, render_template, request, flash, redirect, url_for, session, logging, send_file
import boto3
import datetime, time
import sys, os
from wtforms import Form, StringField, PasswordField, validators, SelectField, SubmitField
from passlib.hash import sha256_crypt
from functools import wraps
import google.cloud.storage as storage
from string import Template
import cloudconvert
import json
import docx
from io import StringIO, BytesIO

app = Flask(__name__)

credential_path = "C:\\Users\\mtryhoen\\OneDrive\\OneDrive - Agilent Technologies\\google\\Speech-1f9b664683c4.json"
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = credential_path

# Check if user logged in
def is_logged_in(f):
    @wraps(f)
    def wrap(*args, **kwargs):
        if 'logged_in' in session:
            return f(*args, **kwargs)
        else:
            flash('Please login', 'danger')
            return redirect(url_for('login'))
    return wrap

#  Session timeout
@app.before_request
def before_request():
    session.permanent = True
    app.permanent_session_lifetime = datetime.timedelta(minutes=30)
    session.modified = True

@app.route('/home')
@is_logged_in
def home():
    return render_template('index.html')

@app.route('/')
@is_logged_in
def index():
    return render_template('index.html')

class TranscriptionForm(Form):
    transcription = StringField('Enter new transcription name', [validators.length(min=4, max=50)])
    template = SelectField(
        'Template',
        coerce=int
    )

@app.route('/transcribe', methods=['GET', 'POST'])
@is_logged_in
def transcription():
    ddb = boto3.client('dynamodb')
    s3 = boto3.client('s3')
    email = session['username']
    form = TranscriptionForm(request.form)
    try:
        objectKey=email.replace('@', '-')
        response = s3.list_objects_v2(
            Bucket="transcribe-template-files",
            Prefix=objectKey + "/templates/"
        )
        templates = response['Contents']
    except:
        templates = []
    templateList = []
    for i in range(len(templates)):
        crap1, crap2, object = str(templates[i]['Key']).split('/')
        if crap2 == "templates":
            data = (str(i), object)
            templateList.append(data)
    form.template.choices = templateList

    try:
        response = ddb.get_item(
            Key={
                'email': {
                    'S': email,
                },
            },
            TableName='users_transcribe',
        )
        transcriptions = response['Item']['transcriptions']['L']
        app.logger.info('Transcriptions: %s', transcriptions)
    except:
        transcriptions = []
    app.logger.info(form.validate())
    if request.method == 'POST' and request.form['btn'] == 'create': #and form.validate():
        transcription = form.transcription.data
        template = form.template.choices[form.template.data]
        app.logger.info('Template: %s', template[1])
        for transcriptionexist in transcriptions:
            if transcription == transcriptionexist['M']['Transcription']['S']:
                flash('transcription already defined', 'danger')
                return redirect(url_for('transcription'))
        try:
            objectKey = email.replace('@', '-')
            url = upload_source_file(request.files.get('sound'), objectKey)
            audiofile = request.files.get('sound').filename
            time.sleep(5)
            if ".mp3" in audiofile:
                print("converting from mp3")
                format = "mp3"
                audioconvert(objectKey + "/" + audiofile, credential_path, format)
                audiofile = audiofile.replace("mp3", "flac")
            elif ".wav" in audiofile:
                format = "wav"
                print("converting from wav")
                audioconvert(objectKey + "/" + audiofile, credential_path, format)
                audiofile = audiofile.replace("wav", "flac")
            elif ".flac" in audiofile:
                print("file already in flac format")
            else:
                print("Unsupported file format")
            gcs_uri="gs://transcribe-sounds/" + objectKey + "/" + audiofile
            app.logger.info('URL: %s', url)
            #return redirect(url_for('transcription'))
        except:
            flash('Could not upload the file', 'danger')
            audiofile = "NA"
            url = "NA"
            return redirect(url_for('transcription'))
        try:
            app.logger.info('Create Transcriptions: %s', response)
            app.logger.info('URI: %s', gcs_uri)
            text = transcribe_gcs(gcs_uri)
            try:
                # open the file
                objectKey = email.replace('@', '-')
                response = s3.get_object(
                    Bucket="transcribe-template-files",
                    Key=objectKey + "/templates/" + template[1]
                )
                if ".docx" in template[1]:
                    try:
                        source_stream = BytesIO(response['Body'].read())
                        document = docx.Document(source_stream)
                        source_stream.close()

                        for para in document.paragraphs:
                            if "$text$" in para.text:
                                cursor_paragraph = para #document.paragraphs[3]
                                print(cursor_paragraph.text)
                                cursor_paragraph.insert_paragraph_before(text=text)
                                delete_paragraph(para)
                                document.save('result.docx')
                                with open('result.docx', 'rb') as data:
                                    response = s3.put_object(
                                        Bucket="transcribe-template-files",
                                        Key=objectKey + "/results/" + transcription + "-" + str(template[1]),
                                        Body=data
                                )
                                break
                        os.remove('result.docx')
                    except:
                        app.logger.info('Error: %s', sys.exc_info()[0])
                elif ".txt" in template[1]:
                    filein = response['Body'].read().decode("utf-8")
                    # read it
                    src = Template(filein)
                    # document data
                    d = {'TEXT': text}
                    # do the substitution
                    result = src.substitute(d)
                    print(result)
                    resultb = str.encode(result)
                    response = s3.put_object(
                        Bucket="transcribe-template-files",
                        Key=objectKey + "/results/" + transcription + "-" + str(template[1]),
                        Body=resultb
                    )

                puburl = s3.generate_presigned_url('get_object', Params={'Bucket': 'transcribe-template-files', 'Key': objectKey + "/results/" + transcription + "-" + str(template[1])},
                                        ExpiresIn=3600)

                response = ddb.update_item(
                    UpdateExpression="SET transcriptions = list_append(transcriptions, :col)",
                    ExpressionAttributeValues={
                        ':col': {
                            "L": [
                                {"M": {"Transcription": {"S": transcription}, "File": {"S": url},
                                       "Template": {"S": template[1]}, "Audiofile": {"S": audiofile},
                                       "Result": {"S": puburl}}}
                            ]
                        },
                    },
                    Key={
                        'email': {
                            'S': email,
                        },
                    },
                    TableName='users_transcribe',
                )

                return redirect(url_for('transcription'))
            except:
                app.logger.info('Error: %s', sys.exc_info()[0])
                return redirect(url_for('transcription'))
        except :
            try:
                response = ddb.update_item(
                    UpdateExpression="SET transcriptions = :col",
                    ExpressionAttributeValues={
                        ':col': {
                            "L": [
                                {"M": {"Transcription": {"S": transcription}, "File": {"S": url},
                                       "Template": {"S": template[1]}, "Audiofile": {"S": audiofile},
                                       "Result": {"S": "Result"}}}
                            ]
                        },
                    },
                    Key={
                        'email': {
                            'S': email,
                        },
                    },
                    TableName='users_transcribe',
                )
                #transcribe_gcs(url)
                return redirect(url_for('transcription'))
            except:
                app.logger.info('Error: %s', sys.exc_info()[0])
                flash('Could not create new transcription', 'danger')
                return redirect(url_for('transcription'))

    elif request.method == 'POST':
        transcriptiondelete = request.form['btn']
        i = 0
        for transcription in transcriptions:
            if transcriptiondelete == transcription['M']['Transcription']['S']:
                try:
                    response = ddb.update_item(
                        UpdateExpression= "REMOVE transcriptions[%(transcription)d]" % {'transcription': i},
                        Key={
                            'email': {
                                'S': email,
                            },
                        },
                        TableName='users_transcribe',
                    )
                    return redirect(url_for('transcription'))
                except:
                    flash('Could not delete transcription', 'danger')
                    return redirect(url_for('transcription'))
            i = i+1
        return render_template('transcriptions.html', form=form, transcriptions=transcriptions)
    else:
        return render_template('transcriptions.html', form=form, transcriptions=transcriptions)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

# def getText(document):
#     fullText = []
#     for para in document.paragraphs:
#         fullText.append(para.text)
#     return str('\n'.join(fullText))

def audioconvert(filename, credential_path, audioformat):
    if credential_path:
        with open(credential_path, 'r') as f:
            gccred = json.load(f)
    api = cloudconvert.Api('')
    try:
        process = api.convert({
            "inputformat": audioformat,
            "outputformat": "flac",
            "input": {
                "googlecloud": {
                    "projectid": "speech-210613",
                    "bucket": "transcribe-sounds",
                    "credentials": {
                        "type": gccred["type"],
                        "project_id": gccred["project_id"],
                        "private_key_id": gccred["private_key_id"],
                        "private_key": gccred["private_key"],
                        "client_email": gccred["client_email"],
                        "client_id": gccred["client_id"],
                        "auth_uri": gccred["auth_uri"],
                        "token_uri": gccred["token_uri"],
                        "auth_provider_x509_cert_url": gccred["auth_provider_x509_cert_url"],
                        "client_x509_cert_url": gccred["client_x509_cert_url"]
                    }
                }
            },
            "file": "" + filename + "",
            "converteroptions": {
                "audio_codec": "FLAC",
                "audio_bitrate": "128",
                "audio_frequency": "8000",
                "strip_metatags": "false"
            },
            "save": True,
            "output": {
                "googlecloud": {
                    "projectid": "speech-210613",
                    "bucket": "transcribe-sounds",
                    "credentials": {
                        "type": gccred["type"],
                        "project_id": gccred["project_id"],
                        "private_key_id": gccred["private_key_id"],
                        "private_key": gccred["private_key"],
                        "client_email": gccred["client_email"],
                        "client_id": gccred["client_id"],
                        "auth_uri": gccred["auth_uri"],
                        "token_uri": gccred["token_uri"],
                        "auth_provider_x509_cert_url": gccred["auth_provider_x509_cert_url"],
                        "client_x509_cert_url": gccred["client_x509_cert_url"]
                    }

                }
            }
        })
    except:
        app.logger.info('Error: %s', sys.exc_info()[0])
    process.wait()

def upload_source_file(file, username):
    """
    Upload the user-uploaded file to Google Cloud Storage and retrieve its
    publicly-accessible URL.
    """
    app.logger.info("file: %s", file)
    if not file:
        app.logger.info("No file")
        return None

    public_url = upload_file(
        file.read(),
        file.filename,
        file.content_type,
        username
    )

    app.logger.info("Uploaded file %s as %s.", file.filename, public_url)

    return public_url

def upload_file(file_stream, filename, content_type, username):
    """
    Uploads a file to a given Cloud Storage bucket and returns the public url
    to the new object.
    """
    app.logger.info("uploading")
    #client = storage.Client(project='speech-210613')
    client = storage.Client() #.from_service_account_json('service_account.json')
    bucket = client.bucket('transcribe-sounds')
    blob = bucket.blob(username + "/" + filename)
    app.logger.info("%s - %s - %s", client, bucket, blob)

    blob.upload_from_string(
        file_stream,
        content_type=content_type,
        client=client
    )

    url = blob.public_url
    app.logger.info("URL %s.", url)

    # if isinstance(url, six.binary_type):
    #     url = url.decode('utf-8')

    return url


def transcribe_gcs(gcs_uri):
    """Asynchronously transcribes the audio file specified by the gcs_uri."""
    from google.cloud import speech
    from google.cloud.speech import enums
    from google.cloud.speech import types
    text = ""
    try:
        app.logger.info("Speech: %s", "YES")
        client = speech.SpeechClient()
        audio = types.RecognitionAudio(uri=gcs_uri)
        config = types.RecognitionConfig(
            encoding=enums.RecognitionConfig.AudioEncoding.FLAC,
            sample_rate_hertz=8000,
            language_code='fr-FR',
            enable_automatic_punctuation=True)

        operation = client.long_running_recognize(config, audio)

        print('Waiting for operation to complete...')
        response = operation.result(timeout=90)

        # Each result is for a consecutive portion of the audio. Iterate through
        # them to get the transcripts for the entire audio file.
        for result in response.results:
            # The first alternative is the most likely one for this portion.
            print(u'Transcript: {}'.format(result.alternatives[0].transcript))
            print('Confidence: {}'.format(result.alternatives[0].confidence))
            text = text + result.alternatives[0].transcript
        return text
    except:
        app.logger.info('Error: %s', sys.exc_info()[0])


def docx_replace(doc_obj, regex , replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text


@app.route('/files')
@is_logged_in
def files():
    return render_template('files.html')


class RegisterForm(Form):
    email = StringField('Email', [validators.length(min=6, max=50)])
    password = PasswordField('Password', [
        validators.DataRequired(),
        validators.EqualTo('confirm', message='passwords do not match')
    ])
    confirm = PasswordField('Confirm Password')

@app.route('/register', methods=['GET', 'POST'])
@is_logged_in
def register():
    form = RegisterForm(request.form)
    ddb = boto3.client('dynamodb')
    email = form.email.data
    emailexist = ''
    try:
        response = ddb.get_item(
            Key={
                'email': {
                    'S': email,
                },
            },
            TableName='users_transcribe',
        )
        emailexist = response['Item']['email']['S']
    except:
        emailexist = ''

    if request.method == 'POST' and form.validate():
        if emailexist != '':
            flash('email address already registered', 'danger')
            return redirect(url_for('register'))
        else:
            password = sha256_crypt.encrypt(str(form.password.data))

            response = ddb.put_item(
                Item={
                    'email': {
                        'S': email,
                    },
                    'password': {
                        'S': password,
                    },
                },
                TableName='users_transcribe',
            )
            flash('You are now registered and can log in', 'success')
            return redirect(url_for('login'))

    return render_template('register.html', form=form)

# user login
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        # get form fields
        email = request.form['email']
        password_candidate = request.form['password']
        ddb = boto3.client('dynamodb')
        try:
            response = ddb.get_item(
                Key={
                    'email': {
                        'S': email,
                    },
                },
                TableName='users_transcribe',
            )

            password = response['Item']['password']['S']

            # compare pwd
            if sha256_crypt.verify(password_candidate, password):
                session['logged_in'] = True
                session['username'] = email

                flash('you are logged in', 'success')
                return redirect(url_for('home'))
            else:
                error = 'Invalid credentials'
                return render_template('login.html', error=error)

        except KeyError:
            error = 'Invalid credentials!'
            return render_template('login.html', error=error)
    else:
        return render_template('login.html')

@app.route('/logout')
@is_logged_in
def logout():
    session.clear()
    flash('you are logged out', 'success')
    return redirect(url_for('login'))


if __name__ == '__main__':
    app.secret_key='TheSecretKey!'
    app.run(host='0.0.0.0', debug=True)